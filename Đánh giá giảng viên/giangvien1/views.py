from flask import render_template, redirect, request, url_for, send_file
from . import gv_bp
from docx import Document
import sqlite3

import os

data = 'sinhvien.db'


@gv_bp.route('/')
def giang_vien():
    conn = sqlite3.connect(data)
    cur = conn.cursor()
    cur.execute("SELECT DISTINCT ten_giang_vien FROM HocPhan")
    gv = cur.fetchall()
    return render_template("giang_vien1.html", data=gv)


@gv_bp.route('/giang_vien/<string:ten_giang_vien>')
def chitiet(ten_giang_vien):
    conn = sqlite3.connect(data)
    conn.row_factory = sqlite3.Row
    cur = conn.cursor()
    cur.execute("SELECT DISTINCT ten_giang_vien FROM HocPhan")
    gv1 = cur.fetchall()
    cur.execute("SELECT * FROM DiemTrungBinh WHERE ten_giang_vien = ?", (ten_giang_vien,))
    gv = cur.fetchall()

    return render_template("diem_gv.html", dataa=gv, data=gv1)


@gv_bp.route('/bao_cao/<string:ten_giang_vien>')
def bao_cao(ten_giang_vien):
    conn = sqlite3.connect(data)
    conn.row_factory = sqlite3.Row
    cur = conn.cursor()

    cur.execute("SELECT * FROM DiemTrungBinh WHERE ten_giang_vien = ?", (ten_giang_vien,))

    dataa = cur.fetchall()

    conn.close()

    return render_template("inbc.html", data=dataa)


@gv_bp.route("/xuat_word/<ten_giang_vien>")
def xuat_word(ten_giang_vien):
    conn = sqlite3.connect(data)
    conn.row_factory = sqlite3.Row
    cur = conn.cursor()

    cur.execute("SELECT * FROM DiemTrungBinh WHERE ten_giang_vien = ?", (ten_giang_vien,))
    dataa = cur.fetchone()

    doc = Document()
    doc.add_heading("BÁO CÁO ĐÁNH GIÁ GIẢNG VIÊN", 0)

    doc.add_paragraph(f"Tên giảng viên: {dataa['ten_giang_vien']}")
    doc.add_paragraph(f"Lớp học phần: {dataa['ma_hoc_phan']}")
    doc.add_paragraph(f"Giảng viên chuẩn bị bài giảng đầy đủ trước khi lên lớp: {dataa['diem_cau_1']}")
    doc.add_paragraph(f"Giảng viên giảng bài rõ ràng và dễ hiểu: {dataa['diem_cau_2']}")
    doc.add_paragraph(f"Giảng viên sẵn sàng giải đáp thắc mắc của sinh viên: {dataa['diem_cau_3']}")
    doc.add_paragraph(f"Phương pháp giảng dạy của giảng viên giúp sinh viên hiểu bài tốt hơn: {dataa['diem_cau_4']}")
    doc.add_paragraph(
        f"Giảng viên tạo môi trường học tập tích cực và khuyến khích sinh viên tham gia thảo luận: {dataa['diem_cau_5']}")
    doc.add_paragraph(f"Điểm trung bình: {dataa['diem_trung_binh']}")

    file_path = os.path.join(os.getcwd(), "bao_cao.docx")
    doc.save(file_path)
    return send_file(file_path, as_attachment=True)
